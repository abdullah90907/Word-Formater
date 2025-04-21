import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to add a border line below the text (no top border for the header)
def add_borders(paragraph, add_top_border=False):
    """Adds borders to the paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # Add top border if required
    if add_top_border:
        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'single')
        top_border.set(qn('w:sz'), '12')  # Border thickness
        top_border.set(qn('w:space'), '1')
        top_border.set(qn('w:color'), '000000')  # Black color
        pBdr.append(top_border)

    # Add bottom border
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '12')  # Border thickness
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')  # Black color
    pBdr.append(bottom_border)

    pPr.append(pBdr)

# Function to apply formatting to text
def apply_formatting(paragraph, font_size, bold=False, alignment=None, no_indent=False):
    """Applies formatting to the paragraph."""
    for run in paragraph.runs:
        font = run.font
        font.name = "Palatino Linotype"
        font.size = Pt(font_size)
        font.color.rgb = RGBColor(0, 0, 0)  # Black color for headings, title, etc
        if bold or run.bold:
            font.bold = True
    paragraph.alignment = alignment if alignment else WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if no_indent:
        paragraph.paragraph_format.left_indent = None
    else:
        paragraph.paragraph_format.left_indent = Inches(0.5)

# Function to adjust images based on width
def adjust_image(image, column_width=3.4, small_dims=(3, 2.1), large_dims=(6.5, 3)):
    if image.width > Inches(column_width):
        image.width, image.height = Inches(large_dims[0]), Inches(large_dims[1])
        image.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        image.width, image.height = Inches(small_dims[0]), Inches(small_dims[1])

# Function to adjust the table to fit inside the column
def adjust_table(table, column_width=3.4):
    for row in table.rows:
        for cell in row.cells:
            # Set the width of each cell to ensure the table fits inside the column width
            cell.width = Inches(column_width)

            # Set cell margins for better spacing
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.paragraph_format.left_indent = Inches(0.01)  # Small indentation for better spacing

    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the table

# Function to format the document
def format_docx(file_path):
    doc = docx.Document(file_path)
    
    # Set up sections and columns
    for section in doc.sections:
        section.different_first_page_header_footer = True
        
        # Set two equal columns
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        
        # Add two equal columns
        columns = section._sectPr.xpath('./w:cols')[0]
        columns.set(qn('w:num'), '2')
        columns.set(qn('w:space'), '120')

        # First page: blank header
        first_page_header = section.first_page_header
        first_page_header.paragraphs.clear()
        
        # Subsequent pages: add header and footer
        if section.header.is_linked_to_previous:
            section.header.is_linked_to_previous = False

        # Add SmartCity to header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = "SmartCity"
        apply_formatting(header_para, font_size=12, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        # Do not add border to header
        add_borders(header_para, add_top_border=False)

        # Add footer with centered page numbers
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), "PAGE")
        footer_para._element.append(fldSimple)

    # Format content
    for para in doc.paragraphs:
        text = para.text.strip().lower()
        
        # Title formatting
        if "title" in para.style.name.lower():
            apply_formatting(para, font_size=18, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER, no_indent=False)
            # Add borders only to title (including bottom and optional top border)
            add_borders(para, add_top_border=True)
        
        # Author formatting
        elif "author" in para.style.name.lower():
            apply_formatting(para, font_size=10, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, no_indent=False)
        
        # Headings formatting
        elif para.style.name.startswith("Heading"):
            apply_formatting(para, font_size=12, bold=True)

        # Content formatting
        else:
            apply_formatting(para, font_size=10, bold=False)

    # Adjust images
    for shape in doc.inline_shapes:
        adjust_image(shape)

    # Adjust tables
    for table in doc.tables:
        adjust_table(table)

    # Save the formatted document
    output_filename = "formatted_document.docx"
    doc.save(output_filename)
    return output_filename

