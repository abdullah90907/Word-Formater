from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
from docx.oxml import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT
import re

# Function to apply formatting
def apply_formatting(paragraph, font_name, font_size, bold=False, italic=False, underline=False, alignment=None):
    """Applies formatting to the paragraph, resetting pre-existing properties."""
    # Reset all relevant paragraph formatting properties to neutral values
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = None  # Reset alignment to default
    paragraph_format.left_indent = Cm(0)  # No left indent
    paragraph_format.right_indent = Cm(0)  # No right indent
    paragraph_format.first_line_indent = Cm(0)  # No first-line indent
    paragraph_format.space_before = Pt(0)  # No space before
    paragraph_format.space_after = Pt(0)  # No space after
    paragraph_format.line_spacing = None  # Reset line spacing to default
    paragraph_format.widow_control = False  # Disable widow/orphan control
    paragraph_format.keep_together = False  # Disable keep together
    paragraph_format.keep_with_next = False  # Disable keep with next

    # Explicitly set alignment if provided
    if alignment is not None:
        paragraph.alignment = alignment

    # Apply run formatting
    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold if bold else run.bold  # Preserve existing bold if not overridden
        font.italic = italic
        font.underline = underline
        font.color.rgb = RGBColor(0, 0, 0)  # Black text

# Function to set page margins and size
def set_page_layout(doc):
    """Sets page margins and size."""
    for section in doc.sections:
        # Set margins to 2.54 cm (1 inch)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Set page size to 21.59 cm (width) and 27.95 cm (height)
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.95)

# Function to add numbering on the right side
def add_numbering(doc):
    """Adds line numbers to the right side of the document."""
    for section in doc.sections:
        sectPr = section._sectPr
        lnNumType = OxmlElement('w:lnNumType')
        lnNumType.set(qn('w:countBy'), '1')
        lnNumType.set(qn('w:distance'), '360')
        lnNumType.set(qn('w:restart'), 'continuous')
        lnNumType.set(qn('w:numStart'), '1')
        sectPr.append(lnNumType)


def clear_header(header):
    """Completely clear all content from a header object."""
    header.is_linked_to_previous = False
    # Remove all child elements from the header
    for child in list(header._element):
        header._element.remove(child)
    # Ensure no residual paragraphs or tables remain
    header.paragraphs.clear()
    header.tables.clear()

def add_header_footer(doc):
    """Applies headers: images on page 1, even/odd headers with continuous numbering from page 2 onward."""
    # Ensure at least two pages for testing (unchanged from original)
    if len(doc.paragraphs) < 2:
        doc.add_page_break()

    # Process each section, ensuring continuous numbering and even/odd headers
    for i, section in enumerate(doc.sections):
        sectPr = section._sectPr

        # Step 1: Clear all headers
        for header_type in ['header', 'first_page_header', 'even_page_header']:
            if hasattr(section, header_type) and getattr(section, header_type) is not None:
                header = getattr(section, header_type)
                clear_header(header)

        # Step 2: Apply headers
        if i == 0:
            # Enable "Different First Page" for section 0
            titlePg = sectPr.find(qn('w:titlePg'))
            if titlePg is None:
                titlePg = OxmlElement('w:titlePg')
                sectPr.append(titlePg)

            # First page header (images only)
            header = section.first_page_header
            header.is_linked_to_previous = False
            table = header.add_table(rows=1, cols=2, width=Cm(16.50))
            table.autofit = False
            table.columns[0].width = Cm(10.795)
            table.columns[1].width = Cm(10.795)
            left_cell = table.cell(0, 0)
            left_paragraph = left_cell.paragraphs[0]
            left_run = left_paragraph.add_run()
            left_run.add_picture("images/left_image.jpg", width=Cm(3))
            left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            right_cell = table.cell(0, 1)
            right_paragraph = right_cell.paragraphs[0]
            right_run = right_paragraph.add_run()
            right_run.add_picture("images/right_image.jpg", width=Cm(3))
            right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Footer (unchanged)
            footer = section.first_page_footer
            footer.is_linked_to_previous = False
            table = footer.add_table(rows=1, cols=2, width=Cm(16.50))
            table.autofit = False
            table.columns[0].width = Cm(5.0)
            table.columns[1].width = Cm(14.50)
            left_cell = table.cell(0, 0)
            left_paragraph = left_cell.paragraphs[0]
            left_run = left_paragraph.add_run()
            left_run.add_picture("images/footer_image.jpg", width=Cm(3))
            left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            right_cell = table.cell(0, 1)
            right_paragraph = right_cell.paragraphs[0]
            right_paragraph.text = (
                "Copyright © 2025 The Author(s). Published by Tech Science Press. "
                "This work is licensed under a Creative Commons Attribution 4.0 International License."
            )
            right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            table.allow_autofit = False

        # Enable Different Odd & Even Pages for all sections
        evenOdd = sectPr.find(qn('w:evenAndOddHeaders'))
        if evenOdd is None:
            evenOdd = OxmlElement('w:evenAndOddHeaders')
            sectPr.append(evenOdd)

        # Even page header (Page number left, Journal right)
        even_header = section.even_page_header
        even_header.is_linked_to_previous = False
        table = even_header.add_table(rows=1, cols=2, width=Cm(16.50))
        table.autofit = False
        table.columns[0].width = Cm(10.795)
        table.columns[1].width = Cm(10.795)
        left_cell = table.cell(0, 0)
        left_paragraph = left_cell.paragraphs[0]
        run = left_paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)
        instr_text = OxmlElement('w:instrText')
        instr_text.text = ' PAGE '
        instr_text.set(qn('xml:space'), 'preserve')
        run._r.append(instr_text)
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)
        left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        right_cell = table.cell(0, 1)
        right_paragraph = right_cell.paragraphs[0]
        right_run = right_paragraph.add_run("Comput Mater Contin. 2025;volume(issue)")
        right_run.font.size = Cm(0.35)  # Match your original size
        right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Odd page header (Journal left, Page number right)
        odd_header = section.header
        odd_header.is_linked_to_previous = False
        table = odd_header.add_table(rows=1, cols=2, width=Cm(16.50))
        table.autofit = False
        table.columns[0].width = Cm(10.795)
        table.columns[1].width = Cm(10.795)
        left_cell = table.cell(0, 0)
        left_paragraph = left_cell.paragraphs[0]
        left_run = left_paragraph.add_run("Comput Mater Contin. 2025;volume(issue)")
        left_run.font.size = Cm(0.35)  # Match your original size
        left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        right_cell = table.cell(0, 1)
        right_paragraph = right_cell.paragraphs[0]
        run = right_paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)
        instr_text = OxmlElement('w:instrText')
        instr_text.text = ' PAGE '
        instr_text.set(qn('xml:space'), 'preserve')
        run._r.append(instr_text)
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)
        right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Step 3: Ensure continuous page numbering
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            sectPr.remove(pgNumType)
        pgNumType = OxmlElement('w:pgNumType')
        if i == 0:
            pgNumType.set(qn('w:start'), '1')  # Start at 1 for the first section only
        # For subsequent sections, omit w:start to continue numbering
        sectPr.append(pgNumType)

# Function to capitalize and bold "Abstract" and "Keyword"
def capitalize_and_bold_abstract_keyword(doc):
    # Define the target words to capitalize and bold
    target_words = ["abstract", "keyword", "keywords"]
    # Iterate through each paragraph
    for paragraph in doc.paragraphs:
        # Check if any target word exists in the paragraph (case-insensitive)
        if any(word in paragraph.text.lower() for word in target_words):
            print(f"Processing paragraph: {paragraph.text}")

            # Clear the paragraph and rebuild it with formatted runs
            new_runs = []
            for run in paragraph.runs:
                text = run.text
                # Capitalize and bold the target words
                for word in target_words:
                    # Use regex to replace all occurrences (case-insensitive)
                    text = re.sub(
                        re.compile(re.escape(word), re.IGNORECASE),
                        word.upper(),  # Capitalize the word
                        text
                    )
                # Add the modified text to new_runs
                new_runs.append((text, run.bold, run.font.size, run.font.name))

            # Clear the paragraph and recreate it with formatted runs
            paragraph.clear()
            for text, orig_bold, orig_size, orig_font in new_runs:
                # Split text into parts (words and separators) using regex
                parts = re.split(r'(\W+)', text)  # Split by non-word characters
                for part in parts:
                    run = paragraph.add_run(part)
                    run.font.name = orig_font or "Minion Pro"
                    run.font.size = orig_size or Pt(10)
                    # Bold only the target words
                    if part.upper() in [word.upper() for word in target_words]:
                        run.bold = True
                    else:
                        run.bold = orig_bold if orig_bold is not None else False


# Function to set a single-column layout
def set_single_column(doc):
    """Ensures the entire document uses a single-column layout."""
    for section in doc.sections:
        section.start_type = WD_SECTION_START.CONTINUOUS  # Prevent new page breaks
        section.page_width = Inches(8.5)  # Standard single-page width
        section.page_height = Inches(11)  # Standard height (A4)

        # Ensure columns are set to 1
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            cols[0].set(qn('w:num'), '1')  # Ensure single-column format

def indent_first_line(paragraph, indent_size=Cm(0.5)):
    """Adds first-line indentation to the paragraph."""
    paragraph.paragraph_format.first_line_indent = indent_size

# Function to identify section based on style
def identify_section(paragraph):
    """Identifies the section type based on the style name from the input document."""
    style_name = paragraph.style.name.lower()

    if "title" in style_name:
        return "title"
    elif "author" in style_name:
        return "author"
    elif "heading 1" in style_name:
        return "heading_1"
    elif "affiliation" in style_name:
        return "affiliation"
    elif "abstract" in style_name:
        return "abstract"
    elif "keyword" in style_name:
        return "keyword"
    elif "articletype" in style_name:
        return "articletype"
    elif "doinum" in style_name:
        return "doinum"
    elif "BackMatter" in style_name:
        return "BackMatter"
    elif "heading 2" in style_name:
        return "heading_2"
    elif "heading 3" in style_name:
        return "heading_3"
    elif "heading 4" in style_name:
        return "heading_4"
    return "body"

def split_and_center_align_images(doc):
    """Splits images into their own paragraphs and center-aligns/resizes them."""
    # Define the namespaces used in the document
    namespaces = {
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    }

    # Iterate through all paragraphs in the document
    for paragraph in list(doc.paragraphs):  # Use list() to avoid skipping paragraphs
        # Check if the paragraph contains both text and images
        has_text = any(run.text.strip() for run in paragraph.runs)
        has_image = any(run.element.xpath('.//w:drawing') for run in paragraph.runs)

        if has_text and has_image:
            # Create a new paragraph for the image
            new_paragraph = doc.add_paragraph()
            new_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Move the image to the new paragraph
            for run in paragraph.runs:
                if run.element.xpath('.//w:drawing'):
                    new_run = new_paragraph.add_run()
                    new_run._r.append(run.element)

            # Remove the image from the original paragraph
            for run in list(paragraph.runs):  # Use list() to avoid skipping runs
                if run.element.xpath('.//w:drawing'):
                    paragraph._element.remove(run._element)

            # Resize the image in the new paragraph
            for run in new_paragraph.runs:
                if run.element.xpath('.//w:drawing'):
                    # Convert the drawing element to an lxml element
                    drawing_xml = run.element.xml
                    drawing = etree.fromstring(drawing_xml)

                    # Resize the image (reduce size to 50% of original)
                    for extent in drawing.xpath('.//wp:extent', namespaces=namespaces):
                        cx = int(int(extent.get('cx')) * 0.85)  # Reduce width by 85%
                        cy = int(int(extent.get('cy')) * 0.85)  # Reduce height by 85%
                        extent.set('cx', str(int(cx)))  # Set new width
                        extent.set('cy', str(int(cy)))  # Set new height

                    # Convert the modified XML back to a string
                    updated_xml = etree.tostring(drawing, encoding='unicode')

                    # Parse the updated XML back into an element
                    updated_element = etree.fromstring(updated_xml)

                    # Update the run's XML with the modified drawing
                    run._r.clear()
                    run._r.append(updated_element)

            # Insert the new paragraph (with the image) immediately after the original paragraph
            paragraph._p.addnext(new_paragraph._element)

def indent_first_line(paragraph, indent_size=Cm(0.5)):
    """Adds first-line indentation to a paragraph."""
    paragraph.paragraph_format.first_line_indent = indent_size

def adjust_table_widths(doc):
    """Adjusts all table widths to the maximum usable width based on set_page_layout."""
    # Use fixed usable width based on set_page_layout: 21.59 cm - 2.54 cm - 2.54 cm = 16.51 cm
    usable_width_cm = Cm(16.51)
    usable_width_twips = int(usable_width_cm.cm * 567)  # Convert to twips (1 cm = 567 twips)

    # Iterate through all tables in the document
    for table in doc.tables:
        # Set table width to maximum usable width
        table.width = usable_width_cm
        table.autofit = False  # Disable autofit to enforce our width

        # Adjust column widths proportionally to fit new table width
        total_current_width = sum(col.width for col in table.columns if col.width is not None)
        if total_current_width > 0:  # Avoid division by zero
            width_ratio = usable_width_cm / total_current_width
            for col in table.columns:
                if col.width is not None:
                    col.width = int(col.width * width_ratio)

        # Set table alignment to left (consistent with your document style)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Ensure table width is enforced via XML
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None:
            tbl_w = OxmlElement('w:tblW')
            tbl_pr.append(tbl_w)
        tbl_w.set(qn('w:w'), str(usable_width_twips))  # Set width in twips
        tbl_w.set(qn('w:type'), 'dxa')  # Use absolute units


def normalize_inline_spacing(doc):
    """Normalizes excessive spaces in all text runs, including headings, preserving images and structure."""
    for para in doc.paragraphs:
        if para._element.findall(qn('w:tbl')) or para._element.xpath('.//w:drawing|.//w:pict'):
            continue
        full_text = ''.join(run.text for run in para.runs if run.text)
        if full_text:
            normalized_text = re.sub(r'\s+', ' ', full_text).strip()
            if normalized_text != full_text:
                para.clear()
                para.add_run(normalized_text)

# Function to detect title, authors, and abstract index
def identify_sections(doc):
    """Identifies title, authors, and abstract index with robust content-based logic."""
    all_paragraphs = [para for para in doc.paragraphs]  # Include empty paragraphs for accurate indexing
    paragraphs = [para.text.strip() for para in all_paragraphs if para.text.strip()]

    # Title: First paragraph that isn’t DOI, Paper Type, or Abstract
    title = ""
    title_index = None
    for i, para in enumerate(all_paragraphs):
        text = para.text.strip()
        if text and not any(text.lower().startswith(x) for x in ["doi", "paper type", "articletype", "abstract"]):
            title = text
            title_index = i
            break

    # Authors: Next non-empty paragraph after title, before Abstract
    authors = ""
    authors_index = None
    if title_index is not None:
        for i in range(title_index + 1, len(all_paragraphs)):
            text = all_paragraphs[i].text.strip()
            if text and not any(text.lower().startswith(x) for x in ["abstract", "doi", "paper type", "articletype"]):
                authors = text
                authors_index = i
                break

    # Abstract index
    abstract_index = next((i for i, para in enumerate(all_paragraphs) if para.text.strip().lower().startswith("abstract")), None)

    return title, authors, title_index, authors_index, abstract_index

def format_docx(file_path):
    """Formats the document and returns the path to the saved file."""
    doc = Document(file_path)

    # Ensure document has at least one paragraph
    if not doc.paragraphs:
        doc.add_paragraph("")

    # Step 1: Handle DOI, Paper Type, and spacing before title
    all_paras = [para for para in doc.paragraphs]
    paragraphs = [para.text.strip() for para in all_paras if para.text.strip()]
    # Stricter DOI detection: must start with "doi:" or "DOI:"
    has_doi = any(p.lower().startswith("doi:") for p in paragraphs if p and len(p) > 4)
    print(f"Has DOI: {has_doi}, Paragraphs checked: {paragraphs}")

    # Insert DOI if absent
    if not has_doi:
        print("Inserting DOI because it’s absent")
        doi_para = doc.add_paragraph("DOI: _________________")
        # Move to top by inserting at the beginning
        if all_paras:  # If there are existing paragraphs
            all_paras[0]._p.addprevious(doi_para._p)
            all_paras.insert(0, doi_para)
        else:  # If document was empty
            all_paras = [doi_para]
        apply_formatting(doi_para, "Minion Pro", 7, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        blank_para = doc.add_paragraph("")
        blank_para.paragraph_format.space_after = Pt(6)
        doi_para._p.addnext(blank_para._p)
        all_paras.insert(1, blank_para)
    else:
        print("DOI already present, formatting only")
        for i, para in enumerate(all_paras):
            if para.text.strip().lower().startswith("doi:"):
                apply_formatting(para, "Minion Pro", 7, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
                break

    # Insert Paper Type if absent
    has_paper_type = any(p.lower().startswith(("paper type", "articletype")) for p in paragraphs if p)
    print(f"Has Paper Type: {has_paper_type}")
    doi_index = next((i for i, para in enumerate(all_paras) if para.text.strip().lower().startswith("doi:")), -1)
    if not has_paper_type:
        print("Inserting Paper Type because it’s absent")
        paper_type_para = doc.add_paragraph("Paper Type (_________________)")
        insert_after = doi_index + 1 if doi_index >= 0 else 0
        if insert_after < len(all_paras) and not all_paras[insert_after].text.strip():
            insert_after += 1  # Skip blank line after DOI
        if insert_after < len(all_paras):
            all_paras[insert_after]._p.addprevious(paper_type_para._p)
            all_paras.insert(insert_after, paper_type_para)
        else:
            doc._body._element.append(paper_type_para._p)
            all_paras.append(paper_type_para)
        apply_formatting(paper_type_para, "Minion Pro", 9, bold=True, underline=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        blank_para = doc.add_paragraph("")
        blank_para.paragraph_format.space_after = Pt(6)
        paper_type_para._p.addnext(blank_para._p)
        all_paras.insert(all_paras.index(paper_type_para) + 1, blank_para)
    else:
        print("Paper Type already present, formatting only")
        for i, para in enumerate(all_paras):
            if para.text.strip().lower().startswith(("paper type", "articletype")):
                apply_formatting(para, "Minion Pro", 9, bold=True, underline=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
                break

    # Identify title and add spacing before it
    title, authors, title_index, authors_index, abstract_index = identify_sections(doc)
    print(f"Title Index: {title_index}, Title: {title}")
    print(f"Authors Index: {authors_index}, Authors: {authors}")
    print(f"Abstract Index: {abstract_index}")

    if title_index is not None and title_index < len(all_paras):
        # Add blank line before title if not already present
        if title_index > 0 and all_paras[title_index - 1].text.strip():
            blank_para = doc.add_paragraph("")
            blank_para.paragraph_format.space_after = Pt(6)
            all_paras[title_index]._p.addprevious(blank_para._p)
            all_paras.insert(title_index, blank_para)
            title_index += 1
            authors_index = authors_index + 1 if authors_index is not None else None
            abstract_index = abstract_index + 1 if abstract_index is not None else None
        elif title_index > 0:
            all_paras[title_index - 1].paragraph_format.space_after = Pt(6)

    # Step 2: Apply title and authors formatting
    if title_index is not None and title_index < len(all_paras):
        apply_formatting(all_paras[title_index], "Minion Pro", 14, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        print(f"Applied title formatting to: {all_paras[title_index].text}")
    if authors_index is not None and authors_index < len(all_paras):
        apply_formatting(all_paras[authors_index], "Minion Pro", 12, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        print(f"Applied authors formatting to: {all_paras[authors_index].text}")
    else:
        print("Authors index not found or invalid!")

    # Ensure single-column layout
    set_single_column(doc)

    # Flags
    in_references_section = False
    before_abstract = True

    # Step 3: Format remaining paragraphs, preserving input spacing
    for i, para in enumerate(all_paras):
        text = para.text.strip()

        # Skip title and authors since they’re already formatted
        if i == title_index or i == authors_index:
            continue

        if abstract_index is not None and i >= abstract_index:
            before_abstract = False

        # Format DOI
        if text.lower().startswith("doi:") and i <= 1:
            apply_formatting(para, "Minion Pro", 7, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
            continue
        # Format Paper Type
        elif text.lower().startswith(("paper type", "articletype")) and i <= 3:
            apply_formatting(para, "Minion Pro", 9, bold=True, underline=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
            continue

        section_type = identify_section(para)
        if section_type == "affiliation":
            apply_formatting(para, "Minion Pro", 9, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif section_type == "abstract" or section_type == "keyword":
            apply_formatting(para, "Minion Pro", 10, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif section_type == "BackMatter":
            apply_formatting(para, "Minion Pro", 10, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif section_type == "heading_1":
            apply_formatting(para, "Minion Pro", 11, bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif section_type == "heading_2":
            apply_formatting(para, "Minion Pro", 11, bold=True, italic=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif section_type == "heading_3" or section_type == "heading_4":
            apply_formatting(para, "Minion Pro", 11, italic=True, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT)
        elif not text:  # Preserve blank lines from input without modification
            continue
        else:
            apply_formatting(para, "Minion Pro", 10, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

        if text.lower().startswith("references"):
            in_references_section = True
            continue
        if not in_references_section and i > 0:
            previous_para = all_paras[i - 1]
            previous_section_type = identify_section(previous_para)
            current_section_type = identify_section(para)
            if previous_section_type in ["heading_1", "heading_2", "heading_3", "heading_4"] and \
               current_section_type not in ["heading_1", "heading_2", "heading_3", "heading_4"]:
                if re.match(r'^[A-Za-z]', text) and not re.match(r'^\d', text) and not re.match(r'^[A-Z]\s', text):
                    if not all(run.bold for run in para.runs if run.text.strip()):
                        indent_first_line(para, Cm(0.5))

    normalize_inline_spacing(doc)
    split_and_center_align_images(doc)
    set_page_layout(doc)
    add_numbering(doc)
    add_header_footer(doc)
    capitalize_and_bold_abstract_keyword(doc)
    adjust_table_widths(doc)

    output_filename = "formated.docx"
    doc.save(output_filename)
    return output_filename
